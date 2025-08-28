# File processing service
# Takes user uploads and turns them into searchable knowledge

import os
import json
import hashlib
import logging
from typing import List, Dict, Optional, Any
from datetime import datetime
import asyncio
import re
from dataclasses import dataclass
import PyPDF2
import docx
import markdown

logger = logging.getLogger(__name__)

@dataclass
class Document:
    """A document we've processed and stored"""
    doc_id: str
    user_id: str
    filename: str
    content: str
    metadata: Dict
    chunks: List[str]
    upload_time: datetime
    doc_type: str

class DataIngestionService:
    def __init__(self):
        """Set up the file processor"""
        self.storage_path = "backend/data/knowledge_base"
        self.max_file_size = 10 * 1024 * 1024  # 10MB
        self.supported_formats = ['.pdf', '.txt', '.docx', '.md', '.json']
        self.chunk_size = 500  # Characters per chunk
        self.documents = {}  # Keep documents in memory for now
        
        # Make sure we have a place to store files
        os.makedirs(self.storage_path, exist_ok=True)
        
        logger.info("Data Ingestion Service initialized")
    
    async def ingest_file(
        self,
        file_path: str,
        user_id: str,
        metadata: Optional[Dict] = None
    ) -> Document:
        """Take a file and add it to the knowledge base"""
        try:
            # Validate file
            if not self._validate_file(file_path):
                raise ValueError(f"Invalid file: {file_path}")
            
            # Extract content based on file type
            content = await self._extract_content(file_path)
            
            # Clean and preprocess content
            cleaned_content = self._clean_content(content)
            
            # Create document chunks for retrieval
            chunks = self._chunk_content(cleaned_content)
            
            # Generate document ID
            doc_id = self._generate_doc_id(user_id, file_path)
            
            # Create document object
            document = Document(
                doc_id=doc_id,
                user_id=user_id,
                filename=os.path.basename(file_path),
                content=cleaned_content,
                metadata=metadata or {},
                chunks=chunks,
                upload_time=datetime.now(),
                doc_type=os.path.splitext(file_path)[1]
            )
            
            # Store document
            self.documents[doc_id] = document
            
            # Save to disk (simple JSON for prototype)
            await self._save_document(document)
            
            logger.info(f"Ingested document {doc_id} for user {user_id}")
            return document
            
        except Exception as e:
            logger.error(f"Failed to ingest file {file_path}: {e}")
            raise
    
    def _validate_file(self, file_path: str) -> bool:
        """Validate file before processing"""
        # Check if file exists
        if not os.path.exists(file_path):
            return False
        
        # Check file size
        if os.path.getsize(file_path) > self.max_file_size:
            logger.warning(f"File too large: {file_path}")
            return False
        
        # Check file extension
        ext = os.path.splitext(file_path)[1].lower()
        if ext not in self.supported_formats:
            logger.warning(f"Unsupported format: {ext}")
            return False
        
        return True
    
    async def _extract_content(self, file_path: str) -> str:
        """Extract text content from file"""
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            
            elif ext == '.pdf':
                return self._extract_pdf_content(file_path)
            
            elif ext == '.docx':
                return self._extract_docx_content(file_path)
            
            elif ext == '.md':
                with open(file_path, 'r', encoding='utf-8') as f:
                    md_content = f.read()
                    # Convert markdown to plain text
                    html = markdown.markdown(md_content)
                    # Simple HTML stripping
                    text = re.sub('<[^<]+?>', '', html)
                    return text
            
            elif ext == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    # Convert JSON to readable text
                    return json.dumps(data, indent=2)
            
        except Exception as e:
            logger.error(f"Content extraction failed: {e}")
            return ""
    
    def _extract_pdf_content(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text.append(page.extract_text())
            return ' '.join(text)
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            return ""
    
    def _extract_docx_content(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return ""
    
    def _clean_content(self, content: str) -> str:
        """Clean and normalize content"""
        # Remove excessive whitespace
        content = re.sub(r'\s+', ' ', content)
        
        # Remove special characters but keep punctuation
        content = re.sub(r'[^\w\s\.\,\!\?\-\:\;]', '', content)
        
        # Trim
        content = content.strip()
        
        return content
    
    def _chunk_content(self, content: str) -> List[str]:
        """Split content into retrievable chunks"""
        chunks = []
        words = content.split()
        
        current_chunk = []
        current_length = 0
        
        for word in words:
            current_chunk.append(word)
            current_length += len(word) + 1  # +1 for space
            
            if current_length >= self.chunk_size:
                chunks.append(' '.join(current_chunk))
                current_chunk = []
                current_length = 0
        
        # Add remaining chunk
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks
    
    def _generate_doc_id(self, user_id: str, file_path: str) -> str:
        """Generate unique document ID"""
        timestamp = datetime.now().isoformat()
        unique_string = f"{user_id}_{file_path}_{timestamp}"
        return hashlib.md5(unique_string.encode()).hexdigest()[:16]
    
    async def _save_document(self, document: Document):
        """Save document to storage"""
        doc_path = os.path.join(self.storage_path, f"{document.doc_id}.json")
        
        doc_data = {
            "doc_id": document.doc_id,
            "user_id": document.user_id,
            "filename": document.filename,
            "content": document.content[:1000],  # Store preview only
            "chunks": document.chunks,
            "metadata": document.metadata,
            "upload_time": document.upload_time.isoformat(),
            "doc_type": document.doc_type
        }
        
        with open(doc_path, 'w') as f:
            json.dump(doc_data, f, indent=2)
    
    async def search_documents(
        self,
        query: str,
        user_id: str,
        limit: int = 5
    ) -> List[Dict[str, Any]]:
        """Search user's documents for relevant content"""
        results = []
        query_lower = query.lower()
        
        # Simple keyword search for prototype
        for doc_id, doc in self.documents.items():
            if doc.user_id != user_id:
                continue
            
            # Search in chunks
            relevant_chunks = []
            for chunk in doc.chunks:
                if query_lower in chunk.lower():
                    score = chunk.lower().count(query_lower)
                    relevant_chunks.append({
                        "chunk": chunk,
                        "score": score
                    })
            
            if relevant_chunks:
                # Sort by relevance
                relevant_chunks.sort(key=lambda x: x["score"], reverse=True)
                
                results.append({
                    "doc_id": doc_id,
                    "filename": doc.filename,
                    "relevant_chunks": relevant_chunks[:3],
                    "upload_time": doc.upload_time.isoformat()
                })
        
        # Sort by number of relevant chunks
        results.sort(key=lambda x: len(x["relevant_chunks"]), reverse=True)
        
        return results[:limit]
    
    async def get_user_documents(self, user_id: str) -> List[Dict]:
        """Get all documents for a user"""
        user_docs = []
        
        for doc_id, doc in self.documents.items():
            if doc.user_id == user_id:
                user_docs.append({
                    "doc_id": doc_id,
                    "filename": doc.filename,
                    "doc_type": doc.doc_type,
                    "upload_time": doc.upload_time.isoformat(),
                    "size": len(doc.content),
                    "chunks": len(doc.chunks)
                })
        
        return user_docs
    
    async def delete_document(self, doc_id: str, user_id: str) -> bool:
        """Delete a document from knowledge base"""
        if doc_id in self.documents:
            doc = self.documents[doc_id]
            
            # Verify ownership
            if doc.user_id != user_id:
                logger.warning(f"Unauthorized delete attempt for doc {doc_id}")
                return False
            
            # Remove from memory
            del self.documents[doc_id]
            
            # Remove from disk
            doc_path = os.path.join(self.storage_path, f"{doc_id}.json")
            if os.path.exists(doc_path):
                os.remove(doc_path)
            
            logger.info(f"Deleted document {doc_id}")
            return True
        
        return False
    
    async def prepare_training_data(self, user_id: str) -> Dict:
        """Prepare user's data for model training"""
        training_data = {
            "user_id": user_id,
            "documents": [],
            "total_content": "",
            "key_topics": [],
            "prepared_at": datetime.now().isoformat()
        }
        
        for doc_id, doc in self.documents.items():
            if doc.user_id == user_id:
                training_data["documents"].append({
                    "filename": doc.filename,
                    "content": doc.content
                })
                training_data["total_content"] += doc.content + " "
        
        # Extract key topics (simple keyword extraction)
        words = training_data["total_content"].lower().split()
        word_freq = {}
        
        for word in words:
            if len(word) > 5:  # Focus on longer words
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Get top topics
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        training_data["key_topics"] = [word for word, freq in sorted_words[:10]]
        
        return training_data