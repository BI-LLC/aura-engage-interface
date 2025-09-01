# backend/app/services/document_processor.py
"""
Document Processing Service for AURA Voice AI
Handles PDF, Word, and text file processing for personal knowledge base
Simple but effective implementation for startup prototype
"""

import os
import io
import hashlib
import json
from typing import List, Dict, Optional, Any
from datetime import datetime
import logging
from dataclasses import dataclass
import asyncio

# Document processing libraries
import PyPDF2
import docx
import markdown
from pathlib import Path

# For text chunking and embeddings
import tiktoken
import openai
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity

logger = logging.getLogger(__name__)

@dataclass
class DocumentChunk:
    """Represents a chunk of document text"""
    chunk_id: str
    document_id: str
    content: str
    metadata: Dict[str, Any]
    embedding: Optional[List[float]] = None
    chunk_index: int = 0

@dataclass
class ProcessedDocument:
    """Represents a processed document"""
    document_id: str
    filename: str
    file_type: str
    content: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    processed_at: str
    total_tokens: int = 0

class DocumentProcessor:
    """
    Handles document upload, processing, and retrieval
    Creates a searchable knowledge base from user documents
    """
    
    def __init__(self):
        """Initialize document processor with simple file storage"""
        # Storage paths
        self.storage_dir = Path("data/documents")
        self.embeddings_dir = Path("data/embeddings")
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        self.embeddings_dir.mkdir(parents=True, exist_ok=True)
        
        # OpenAI for embeddings (reuse existing key)
        from app.config import settings
        self.openai_key = settings.OPENAI_API_KEY
        if self.openai_key:
            openai.api_key = self.openai_key
        
        # Token counter for OpenAI
        self.tokenizer = tiktoken.encoding_for_model("gpt-4")
        
        # In-memory document store (for prototype - use vector DB in production)
        self.documents = {}  # document_id -> ProcessedDocument
        self.embeddings_cache = {}  # chunk_id -> embedding vector
        
        # Configuration
        self.chunk_size = 500  # tokens per chunk
        self.chunk_overlap = 50  # token overlap between chunks
        self.max_file_size = 10 * 1024 * 1024  # 10MB max
        
        logger.info("Document Processor initialized")
        self._load_existing_documents()
    
    def _load_existing_documents(self):
        """Load previously processed documents from disk"""
        try:
            # Load document metadata
            metadata_file = self.storage_dir / "documents.json"
            if metadata_file.exists():
                with open(metadata_file, 'r') as f:
                    stored_docs = json.load(f)
                    for doc_id, doc_data in stored_docs.items():
                        # Reconstruct document objects
                        chunks = [DocumentChunk(**chunk) for chunk in doc_data['chunks']]
                        doc = ProcessedDocument(**{**doc_data, 'chunks': chunks})
                        self.documents[doc_id] = doc
                
                logger.info(f"Loaded {len(self.documents)} existing documents")
            
            # Load embeddings
            embeddings_file = self.embeddings_dir / "embeddings.npz"
            if embeddings_file.exists():
                data = np.load(embeddings_file, allow_pickle=True)
                self.embeddings_cache = dict(data['embeddings'].item())
                logger.info(f"Loaded {len(self.embeddings_cache)} embeddings")
                
        except Exception as e:
            logger.error(f"Error loading existing documents: {e}")
    
    def _save_documents(self):
        """Save processed documents to disk for persistence"""
        try:
            # Save document metadata
            metadata_file = self.storage_dir / "documents.json"
            stored_docs = {}
            
            for doc_id, doc in self.documents.items():
                doc_dict = {
                    'document_id': doc.document_id,
                    'filename': doc.filename,
                    'file_type': doc.file_type,
                    'content': doc.content[:1000],  # Save preview only
                    'chunks': [
                        {
                            'chunk_id': chunk.chunk_id,
                            'document_id': chunk.document_id,
                            'content': chunk.content,
                            'metadata': chunk.metadata,
                            'chunk_index': chunk.chunk_index
                        }
                        for chunk in doc.chunks
                    ],
                    'metadata': doc.metadata,
                    'processed_at': doc.processed_at,
                    'total_tokens': doc.total_tokens
                }
                stored_docs[doc_id] = doc_dict
            
            with open(metadata_file, 'w') as f:
                json.dump(stored_docs, f, indent=2)
            
            # Save embeddings
            if self.embeddings_cache:
                embeddings_file = self.embeddings_dir / "embeddings.npz"
                np.savez_compressed(embeddings_file, embeddings={'embeddings': self.embeddings_cache})
            
            logger.info(f"Saved {len(self.documents)} documents to disk")
            
        except Exception as e:
            logger.error(f"Error saving documents: {e}")
    
    async def process_document(self, file_content: bytes, filename: str, user_id: str) -> ProcessedDocument:
        """
        Process a document and add to knowledge base
        Supports PDF, DOCX, TXT, MD files
        """
        try:
            # Check file size
            if len(file_content) > self.max_file_size:
                raise ValueError(f"File too large. Max size is {self.max_file_size / 1024 / 1024}MB")
            
            # Generate document ID
            doc_id = hashlib.md5(f"{filename}:{user_id}:{datetime.now().isoformat()}".encode()).hexdigest()[:12]
            
            # Extract text based on file type
            file_ext = filename.lower().split('.')[-1]
            
            if file_ext == 'pdf':
                text = self._extract_pdf_text(file_content)
                file_type = 'pdf'
            elif file_ext in ['docx', 'doc']:
                text = self._extract_docx_text(file_content)
                file_type = 'docx'
            elif file_ext == 'md':
                text = self._extract_markdown_text(file_content)
                file_type = 'markdown'
            elif file_ext == 'txt':
                text = file_content.decode('utf-8', errors='ignore')
                file_type = 'text'
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            if not text or len(text.strip()) < 10:
                raise ValueError("Document contains no extractable text")
            
            # Clean and normalize text
            text = self._clean_text(text)
            
            # Create chunks for better retrieval
            chunks = self._create_chunks(text, doc_id)
            
            # Generate embeddings for each chunk
            await self._generate_embeddings(chunks)
            
            # Count total tokens
            total_tokens = sum(len(self.tokenizer.encode(chunk.content)) for chunk in chunks)
            
            # Create document object
            document = ProcessedDocument(
                document_id=doc_id,
                filename=filename,
                file_type=file_type,
                content=text[:5000],  # Store preview
                chunks=chunks,
                metadata={
                    'user_id': user_id,
                    'file_size': len(file_content),
                    'chunk_count': len(chunks),
                    'original_filename': filename
                },
                processed_at=datetime.now().isoformat(),
                total_tokens=total_tokens
            )
            
            # Store document
            self.documents[doc_id] = document
            
            # Save to disk
            self._save_documents()
            
            logger.info(f"Processed document: {filename} ({len(chunks)} chunks, {total_tokens} tokens)")
            
            return document
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            raise
    
    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n\n"
            
            return text
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise ValueError(f"Failed to extract text from PDF: {e}")
    
    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from Word document"""
        try:
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return "\n\n".join(text)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise ValueError(f"Failed to extract text from Word document: {e}")
    
    def _extract_markdown_text(self, file_content: bytes) -> str:
        """Extract text from Markdown file"""
        try:
            md_text = file_content.decode('utf-8', errors='ignore')
            # Convert markdown to plain text (remove formatting)
            html = markdown.markdown(md_text)
            # Simple HTML tag removal
            import re
            text = re.sub('<[^<]+?>', '', html)
            return text
        except Exception as e:
            logger.error(f"Markdown extraction error: {e}")
            raise ValueError(f"Failed to extract text from Markdown: {e}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove special characters but keep punctuation
        text = ''.join(char if char.isprintable() else ' ' for char in text)
        
        # Ensure proper spacing
        text = text.replace('  ', ' ')
        
        return text.strip()
    
    def _create_chunks(self, text: str, doc_id: str) -> List[DocumentChunk]:
        """
        Split text into chunks for better retrieval
        Uses token-based chunking with overlap
        """
        chunks = []
        tokens = self.tokenizer.encode(text)
        
        # Calculate chunk boundaries
        chunk_index = 0
        start_idx = 0
        
        while start_idx < len(tokens):
            # Get chunk tokens
            end_idx = min(start_idx + self.chunk_size, len(tokens))
            chunk_tokens = tokens[start_idx:end_idx]
            
            # Decode back to text
            chunk_text = self.tokenizer.decode(chunk_tokens)
            
            # Create chunk object
            chunk_id = f"{doc_id}_chunk_{chunk_index}"
            chunk = DocumentChunk(
                chunk_id=chunk_id,
                document_id=doc_id,
                content=chunk_text,
                metadata={
                    'start_token': start_idx,
                    'end_token': end_idx,
                    'chunk_size': len(chunk_tokens)
                },
                chunk_index=chunk_index
            )
            
            chunks.append(chunk)
            
            # Move to next chunk with overlap
            start_idx = end_idx - self.chunk_overlap
            chunk_index += 1
            
            # Prevent infinite loop
            if start_idx >= len(tokens) - 1:
                break
        
        logger.info(f"Created {len(chunks)} chunks from text")
        return chunks
    
    async def _generate_embeddings(self, chunks: List[DocumentChunk]):
        """Generate embeddings for document chunks using OpenAI"""
        if not self.openai_key:
            logger.warning("OpenAI API key not configured - skipping embeddings")
            return
        
        try:
            client = openai.AsyncOpenAI(api_key=self.openai_key)
            
            for chunk in chunks:
                # Check if embedding already exists
                if chunk.chunk_id in self.embeddings_cache:
                    chunk.embedding = self.embeddings_cache[chunk.chunk_id]
                    continue
                
                # Generate embedding
                response = await client.embeddings.create(
                    model="text-embedding-ada-002",
                    input=chunk.content[:8000]  # Max tokens for embedding model
                )
                
                embedding = response.data[0].embedding
                chunk.embedding = embedding
                
                # Cache embedding
                self.embeddings_cache[chunk.chunk_id] = embedding
                
                # Small delay to avoid rate limiting
                await asyncio.sleep(0.1)
            
            logger.info(f"Generated embeddings for {len(chunks)} chunks")
            
        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            # Continue without embeddings - fallback to keyword search
    
    async def search_documents(self, query: str, user_id: str, top_k: int = 3) -> List[Dict[str, Any]]:
        """
        Search documents for relevant information
        Uses semantic search with embeddings if available, falls back to keyword search
        """
        try:
            # Filter documents by user
            user_docs = [doc for doc in self.documents.values() 
                        if doc.metadata.get('user_id') == user_id]
            
            if not user_docs:
                return []
            
            # Try semantic search if embeddings are available
            if self.openai_key and self.embeddings_cache:
                return await self._semantic_search(query, user_docs, top_k)
            else:
                # Fallback to keyword search
                return self._keyword_search(query, user_docs, top_k)
            
        except Exception as e:
            logger.error(f"Document search error: {e}")
            return []
    
    async def _semantic_search(self, query: str, documents: List[ProcessedDocument], top_k: int) -> List[Dict[str, Any]]:
        """Semantic search using embeddings"""
        try:
            # Generate query embedding
            client = openai.AsyncOpenAI(api_key=self.openai_key)
            response = await client.embeddings.create(
                model="text-embedding-ada-002",
                input=query
            )
            query_embedding = np.array(response.data[0].embedding)
            
            # Calculate similarities
            results = []
            
            for doc in documents:
                for chunk in doc.chunks:
                    if chunk.chunk_id in self.embeddings_cache:
                        chunk_embedding = np.array(self.embeddings_cache[chunk.chunk_id])
                        
                        # Calculate cosine similarity
                        similarity = cosine_similarity(
                            query_embedding.reshape(1, -1),
                            chunk_embedding.reshape(1, -1)
                        )[0][0]
                        
                        results.append({
                            'content': chunk.content,
                            'score': float(similarity),
                            'document': doc.filename,
                            'chunk_index': chunk.chunk_index,
                            'metadata': chunk.metadata
                        })
            
            # Sort by similarity and return top k
            results.sort(key=lambda x: x['score'], reverse=True)
            return results[:top_k]
            
        except Exception as e:
            logger.error(f"Semantic search error: {e}")
            # Fallback to keyword search
            return self._keyword_search(query, documents, top_k)
    
    def _keyword_search(self, query: str, documents: List[ProcessedDocument], top_k: int) -> List[Dict[str, Any]]:
        """Simple keyword-based search fallback"""
        query_words = set(query.lower().split())
        results = []
        
        for doc in documents:
            for chunk in doc.chunks:
                # Count matching words
                chunk_words = set(chunk.content.lower().split())
                matches = len(query_words.intersection(chunk_words))
                
                if matches > 0:
                    results.append({
                        'content': chunk.content,
                        'score': matches / len(query_words),  # Simple relevance score
                        'document': doc.filename,
                        'chunk_index': chunk.chunk_index,
                        'metadata': chunk.metadata
                    })
        
        # Sort by score
        results.sort(key=lambda x: x['score'], reverse=True)
        return results[:top_k]
    
    def get_user_documents(self, user_id: str) -> List[Dict[str, Any]]:
        """Get list of documents for a user"""
        user_docs = []
        
        for doc in self.documents.values():
            if doc.metadata.get('user_id') == user_id:
                user_docs.append({
                    'document_id': doc.document_id,
                    'filename': doc.filename,
                    'file_type': doc.file_type,
                    'processed_at': doc.processed_at,
                    'chunk_count': len(doc.chunks),
                    'total_tokens': doc.total_tokens
                })
        
        return user_docs
    
    def delete_document(self, document_id: str, user_id: str) -> bool:
        """Delete a document from the knowledge base"""
        try:
            if document_id in self.documents:
                doc = self.documents[document_id]
                
                # Verify user owns the document
                if doc.metadata.get('user_id') != user_id:
                    logger.warning(f"User {user_id} attempted to delete document owned by another user")
                    return False
                
                # Remove embeddings
                for chunk in doc.chunks:
                    if chunk.chunk_id in self.embeddings_cache:
                        del self.embeddings_cache[chunk.chunk_id]
                
                # Remove document
                del self.documents[document_id]
                
                # Save changes
                self._save_documents()
                
                logger.info(f"Deleted document {document_id} for user {user_id}")
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            return False
    
    def get_context_for_chat(self, query: str, user_id: str) -> str:
        """
        Get relevant context from documents for chat
        Returns formatted context string to inject into LLM prompt
        """
        try:
            # Search documents
            results = asyncio.run(self.search_documents(query, user_id, top_k=2))
            
            if not results:
                return ""
            
            # Format context
            context_parts = []
            for i, result in enumerate(results, 1):
                context_parts.append(
                    f"[Document: {result['document']}]\n{result['content'][:500]}"
                )
            
            context = "\n\n".join(context_parts)
            
            # Create context prompt
            return f"""Based on your uploaded documents, here's relevant information:

{context}

Using this context, I'll answer your question."""
            
        except Exception as e:
            logger.error(f"Error getting chat context: {e}")
            return ""